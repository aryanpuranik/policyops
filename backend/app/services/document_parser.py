import os
import io
import aiofiles
from pathlib import Path
from typing import Optional


async def parse_document(file_path: str, file_type: str) -> str:
    """Parse a document and extract its text content."""
    file_type = file_type.lower().strip(".")

    if file_type in ("pdf",):
        return await _parse_pdf(file_path)
    elif file_type in ("docx", "doc"):
        return await _parse_docx(file_path)
    elif file_type in ("txt", "text", "md"):
        return await _parse_text(file_path)
    else:
        return await _parse_text(file_path)


async def _parse_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfplumber."""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"[Page {page_num}]\n{page_text}")

            # Extract tables as structured text
            tables = page.extract_tables()
            for table in tables:
                if table:
                    table_text = _table_to_text(table)
                    if table_text:
                        text_parts.append(f"[Table on Page {page_num}]\n{table_text}")

    return "\n\n".join(text_parts)


async def _parse_docx(file_path: str) -> str:
    """Extract text from Word document."""
    from docx import Document

    doc = Document(file_path)
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            style = para.style.name if para.style else "Normal"
            if "Heading" in style:
                text_parts.append(f"\n## {para.text}")
            else:
                text_parts.append(para.text)

    for table in doc.tables:
        text_parts.append("\n[Table]")
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                text_parts.append(row_text)

    return "\n".join(text_parts)


async def _parse_text(file_path: str) -> str:
    """Read plain text file."""
    async with aiofiles.open(file_path, mode="r", encoding="utf-8", errors="ignore") as f:
        return await f.read()


def _table_to_text(table: list) -> str:
    """Convert a table (list of lists) to readable text."""
    if not table:
        return ""
    rows = []
    for row in table:
        if row:
            row_text = " | ".join(str(cell or "").strip() for cell in row)
            if row_text.strip(" |"):
                rows.append(row_text)
    return "\n".join(rows)


async def save_upload(file_content: bytes, filename: str, upload_dir: str) -> str:
    """Save uploaded file to disk and return path."""
    Path(upload_dir).mkdir(parents=True, exist_ok=True)
    safe_name = "".join(c if c.isalnum() or c in "._- " else "_" for c in filename)
    file_path = os.path.join(upload_dir, safe_name)

    async with aiofiles.open(file_path, "wb") as f:
        await f.write(file_content)

    return file_path
